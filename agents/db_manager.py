"""
Agent 3: Database Manager
Saves extracted expressions to:
  1. Master Excel DB (01_Database/english_expressions_db.xlsx)
  2. Daily Excel Sheet (02_Daily_Sheets/Expressions_YYYY-MM-DD.xlsx)
  3. Study Markdown Note (03_Print_PDF/Study_Note_YYYY-MM-DD.md)
"""

import os
import json
import datetime
import openpyxl
from utils.logger import setup_logger
from utils.dedup import get_next_uid

import config

logger = setup_logger('db_manager')

HEADERS = [
    'UID', 'Date', 'Source', 'Expression', 'POS',
    'Pronunciation', 'Meaning_KR', 'Original_Text', 'Applied_Example'
]


def _generate_uid(date_str: str, sequence_num: int) -> str:
    """Generate a unique identifier for an expression entry: ENG-YYYYMMDD-NNN"""
    date_compact = date_str.replace('-', '')
    return f"ENG-{date_compact}-{sequence_num:03d}"


def _prepare_rows(expressions: list[dict], date_str: str, index_data: dict) -> list[list]:
    """Convert expression dicts into row lists suitable for spreadsheet insertion."""
    rows = []

    for expr in expressions:
        seq_num = get_next_uid(index_data, date_str)
        uid = _generate_uid(date_str, seq_num)
        
        # Attach the UID back to the expression dictionary for markdown generation
        expr['uid'] = uid

        row = [
            uid,
            date_str,
            expr.get('source', 'Inspiration'),
            expr.get('expression', ''),
            expr.get('pos', ''),
            expr.get('ipa', ''),
            expr.get('meaning_kr', ''),
            expr.get('original_text', ''),
            expr.get('applied_example', '')
        ]
        rows.append(row)

    return rows


def _save_master_excel(rows: list[list], headers: list[str]) -> int:
    """Save expression rows to the master Excel file (01_Database/english_expressions_db.xlsx)."""
    file_path = config.EXCEL_FILENAME
    os.makedirs(config.DB_DIR, exist_ok=True)

    try:
        if os.path.exists(file_path):
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active
            logger.info(f"Opened master Excel: {file_path} (last row: {ws.max_row})")
        else:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Expressions"
            ws.append(headers)
            logger.info(f"Created new master Excel: {file_path}")

        for row in rows:
            ws.append(row)

        wb.save(file_path)
        logger.info(f"Appended {len(rows)} expressions to master Excel.")
        return len(rows)

    except Exception as e:
        logger.error(f"Failed to save master Excel: {e}")
        raise


def _save_daily_excel(rows: list[list], headers: list[str], date_str: str) -> None:
    """Save daily expression rows to a dedicated sheet (02_Daily_Sheets/Expressions_YYYY-MM-DD.xlsx)."""
    os.makedirs(config.DAILY_DIR, exist_ok=True)
    file_path = os.path.join(config.DAILY_DIR, f"Expressions_{date_str}.xlsx")

    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = f"Daily_{date_str}"
        ws.append(headers)

        for row in rows:
            ws.append(row)

        wb.save(file_path)
        logger.info(f"Saved daily Excel sheet: {file_path}")

    except Exception as e:
        logger.error(f"Failed to save daily Excel: {e}")
        raise


def _save_word_note(expressions: list[dict], date_str: str) -> None:
    """Save daily expressions to a Word (.docx) document with card-style layout.
    Each expression is a self-contained card: headword + pronunciation + meaning + context + example.
    No tables, no pipe characters.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    os.makedirs(config.PRINT_DIR, exist_ok=True)
    file_path = os.path.join(config.PRINT_DIR, f"Study_Note_{date_str}.docx")

    try:
        doc = Document()

        # ── 페이지 여백 설정 (좌우 2cm, 상하 1.8cm) ─────────────────────────
        section = doc.sections[0]
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

        # ── 기본 폰트 설정 헬퍼 ──────────────────────────────────────────────
        def _set_font(run, name="맑은 고딕", size=11, bold=False, italic=False, color=None):
            run.font.name = name
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)
            # 한글 폰트도 같이 지정
            r = run._r
            rPr = r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), name)
            rPr.append(rFonts)

        def _para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            if text:
                p.add_run(text)
            return p

        # ── 타이틀 ───────────────────────────────────────────────────────────
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after  = Pt(6)
        title_r = title_p.add_run(f"📝 Daily English Study Note  |  {date_str}")
        _set_font(title_r, size=15, bold=True, color=(30, 100, 200))

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_before = Pt(0)
        sub_p.paragraph_format.space_after  = Pt(14)
        sub_r = sub_p.add_run(f"오늘의 핵심 네이티브 표현  {len(expressions)}개  ·  발음 기호 · 의미 · 원문 · 실전 예문 수록")
        _set_font(sub_r, size=9, italic=True, color=(100, 100, 100))

        # ── 표현 카드 반복 ────────────────────────────────────────────────────
        for i, expr in enumerate(expressions, 1):
            expression = expr.get('expression', '')
            pos        = expr.get('pos', '').upper()
            ipa        = expr.get('ipa', '')
            meaning    = expr.get('meaning_kr', '')
            orig_txt   = expr.get('original_text', '')
            appl_ex    = expr.get('applied_example', '')
            source     = expr.get('source', '')

            # 카드 번호 + 표제어 + 품사
            head_p = doc.add_paragraph()
            head_p.paragraph_format.space_before = Pt(10)
            head_p.paragraph_format.space_after  = Pt(2)

            num_r = head_p.add_run(f"{i:03d}.  ")
            _set_font(num_r, size=10, bold=False, color=(150, 150, 150))

            expr_r = head_p.add_run(expression)
            _set_font(expr_r, name="Calibri", size=14, bold=True, color=(20, 60, 180))

            pos_r = head_p.add_run(f"   [{pos}]")
            _set_font(pos_r, name="Calibri", size=9, bold=False, color=(120, 120, 120))

            # 발음 기호
            ipa_p = doc.add_paragraph()
            ipa_p.paragraph_format.space_before = Pt(0)
            ipa_p.paragraph_format.space_after  = Pt(2)
            ipa_label = ipa_p.add_run("🔊 ")
            _set_font(ipa_label, size=10)
            ipa_r = ipa_p.add_run(ipa)
            _set_font(ipa_r, name="Calibri", size=10, italic=True, color=(80, 80, 80))

            # 한국어 의미
            mean_p = doc.add_paragraph()
            mean_p.paragraph_format.space_before = Pt(0)
            mean_p.paragraph_format.space_after  = Pt(3)
            mean_label = mean_p.add_run("💡 의미  ")
            _set_font(mean_label, size=10, bold=True, color=(40, 40, 40))
            mean_r = mean_p.add_run(meaning)
            _set_font(mean_r, size=10.5)

            # 원문 Context
            orig_p = doc.add_paragraph()
            orig_p.paragraph_format.space_before = Pt(0)
            orig_p.paragraph_format.space_after  = Pt(2)
            orig_label = orig_p.add_run("📌 원문  ")
            _set_font(orig_label, size=9.5, bold=True, color=(100, 100, 100))
            orig_r = orig_p.add_run(f'"{orig_txt}"')
            _set_font(orig_r, name="Calibri", size=9.5, italic=True, color=(90, 90, 90))

            # 실전 예문
            ex_p = doc.add_paragraph()
            ex_p.paragraph_format.space_before = Pt(2)
            ex_p.paragraph_format.space_after  = Pt(2)
            ex_label = ex_p.add_run("✏️ 예문  ")
            _set_font(ex_label, size=10, bold=True, color=(0, 120, 80))
            ex_r = ex_p.add_run(appl_ex)
            _set_font(ex_r, name="Calibri", size=10.5, bold=True, color=(0, 100, 60))

            # 구분선 (마지막 카드 제외)
            if i < len(expressions):
                sep_p = doc.add_paragraph()
                sep_p.paragraph_format.space_before = Pt(6)
                sep_p.paragraph_format.space_after  = Pt(0)
                sep_r = sep_p.add_run("─" * 60)
                _set_font(sep_r, size=7, color=(200, 200, 200))

        doc.save(file_path)
        logger.info(f"Saved daily study Word note: {file_path}")

    except Exception as e:
        logger.error(f"Failed to save daily Word note: {e}")
        raise


def save_expressions(expressions: list[dict], index_data: dict) -> int:
    """
    Main entry point for saving expressions.
    Saves to master Excel DB, daily Excel sheets, and markdown notes.
    """
    if not expressions:
        logger.warning("No expressions to save")
        return 0

    date_str = datetime.date.today().strftime('%Y-%m-%d')

    # 1. Convert expressions list of dicts to rows lists
    rows = _prepare_rows(expressions, date_str, index_data)
    logger.info(f"Prepared {len(rows)} rows for saving (date: {date_str})")

    # 2. Save master Excel database (append mode)
    saved_count = _save_master_excel(rows, HEADERS)

    # 3. Save daily-only Excel sheet (new file)
    _save_daily_excel(rows, HEADERS, date_str)

    # 4. Save daily Word study note (new file)
    _save_word_note(expressions, date_str)

    return saved_count
